"""Word 文档工具 — 读取 .docx 文件中的文本和表格。

模块导入时通过 ``registry.register()`` 注册。
依赖 ``python-docx``（请参见 requirements.txt）。
"""

from __future__ import annotations

import logging
from typing import Any, Dict, List

from abstract.tools.registry import registry, tool_error, tool_result
from component.tools.filesystem import _s as _get_sandbox

logger = logging.getLogger(__name__)

_docx = None
try:
    import docx as _docx
except ImportError:
    pass


# ---------------------------------------------------------------------------
# Handlers
# ---------------------------------------------------------------------------


def _handle_read_docx(args: Dict[str, Any]) -> dict:
    path: str = str(args.get("path", "")).strip()
    if not path:
        return tool_error("path is required")

    if _docx is None:
        return tool_error(
            "python-docx is not installed. Run: pip install python-docx",
        )

    try:
        sandbox = _get_sandbox()
        r = sandbox.resolve_read(path)
    except Exception as e:
        return tool_error(str(e), path=path)

    try:
        doc = _docx.Document(str(r.real))
    except Exception as e:
        return tool_error(f"Failed to open document: {e}", path=path)

    # 段落文本
    paragraphs: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                "style": para.style.name if para.style else "Normal",
                "text": text,
            })

    # 表格
    tables: list[dict[str, Any]] = []
    for table in doc.tables:
        rows_data: list[dict[str, str]] = []
        headers: list[str] = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = [c or f"col_{i}" for i, c in enumerate(cells)]
            else:
                row_dict: dict[str, str] = {}
                for i, cell in enumerate(cells):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    row_dict[key] = cell
                if any(row_dict.values()):
                    rows_data.append(row_dict)

        tables.append({
            "headers": headers,
            "rows": rows_data,
            "row_count": len(rows_data),
        })

    # 纯文本（用于快速摘要）
    plain_text: str = "\n".join(
        p["text"] for p in paragraphs
    )

    return tool_result(
        path=path,
        paragraphs=paragraphs,
        tables=tables,
        plain_text=plain_text,
        paragraph_count=len(paragraphs),
        table_count=len(tables),
    )


# ---------------------------------------------------------------------------
# Registration
# ---------------------------------------------------------------------------

registry.register(
    name="read_docx",
    toolset="extools",
    schema={
        "description": (
            # 读取 .docx Word 文档。返回：
            # 1) paragraphs — 段落列表（含文本和样式名称）
            # 2) tables — 文档中的表格（含表头和数据行）
            # 3) plain_text — 纯文本拼接，便于快速预览
            # 空段落会被跳过。
            "Read a .docx Word document. Returns:\n"
            "1) paragraphs — list of paragraphs (with text and style name)\n"
            "2) tables — tables in the document (with headers and data rows)\n"
            "3) plain_text — concatenated plain text for quick preview\n"
            "Empty paragraphs are skipped."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    # Word 文档逻辑路径，必须使用命名空间前缀（ws:、fork:）。例如 'ws:docs/report.docx'。
                    "description": (
                        "Word document logical path, must use a namespace prefix "
                        "(ws:, fork:). E.g. 'ws:docs/report.docx'."
                    ),
                },
            },
            "required": ["path"],
        },
    },
    handler=_handle_read_docx,
    emoji="📄",
)