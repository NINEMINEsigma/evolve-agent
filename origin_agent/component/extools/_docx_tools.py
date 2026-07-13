"""Word 文档工具 — 读取 .docx 文件中的文本和表格。

模块导入时通过 ``registry.register()`` 注册。
依赖 ``python-docx``（请参见 requirements.txt）。
"""

from __future__ import annotations

import logging
from typing import Any, Dict, List

from abstract.tools.registry import registry, tool_error, tool_result
from component.tools.filesystem import _s as _get_sandbox

logger = logging.getLogger(__name__)

_docx = None
try:
    import docx as _docx
except ImportError:
    pass


# ---------------------------------------------------------------------------
# Handlers
# ---------------------------------------------------------------------------


def _handle_read_docx(args: dict[str, Any]) -> dict:
    path: str = str(args.get("path", "")).strip()
    if not path:
        return tool_error("path is required")

    if _docx is None:
        return tool_error(
            "python-docx is not installed. Run: pip install python-docx",
        )

    try:
        sandbox = _get_sandbox()
        r = sandbox.resolve_read(path)
    except Exception as e:
        return tool_error(str(e), path=path)

    try:
        doc = _docx.Document(str(r.real))
    except Exception as e:
        return tool_error(f"Failed to open document: {e}", path=path)

    # 段落文本
    paragraphs: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                "style": para.style.name if para.style else "Normal",
                "text": text,
            })

    # 表格
    tables: list[dict[str, Any]] = []
    for table in doc.tables:
        rows_data: list[dict[str, str]] = []
        headers: list[str] = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = [c or f"col_{i}" for i, c in enumerate(cells)]
            else:
                row_dict: dict[str, str] = {}
                for i, cell in enumerate(cells):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    row_dict[key] = cell
                if any(row_dict.values()):
                    rows_data.append(row_dict)

        tables.append({
            "headers": headers,
            "rows": rows_data,
            "row_count": len(rows_data),
        })

    # 纯文本（用于快速摘要）
    plain_text: str = "\n".join(
        p["text"] for p in paragraphs
    )

    return tool_result(
        path=path,
        paragraphs=paragraphs,
        tables=tables,
        plain_text=plain_text,
        paragraph_count=len(paragraphs),
        table_count=len(tables),
    )


# ---------------------------------------------------------------------------
# Registration
# ---------------------------------------------------------------------------

registry.register(
    name="read_docx",
    toolset="extools",
    schema={
        # 读取 .docx Word 文档。
        #
        # ## 前置条件
        # 必须安装 python-docx。
        # path 必须使用命名空间前缀（如 ws:、fork:）。
        #
        # ## 调用效果
        # 返回段落列表（含文本和样式名）、表格列表（含表头和数据行）以及拼接后的纯文本摘要。
        # 空段落会被跳过。
        #
        # ## 返回
        # ```json
        # {"path": "ws:docs/report.docx", "paragraphs": [{"style": "Normal", "text": "..."}], "tables": [{"headers": ["..."], "rows": [{"...": "..."}], "row_count": 1}], "plain_text": "...", "paragraph_count": 10, "table_count": 2}
        # ```
        #
        # ## 何时使用
        # - 读取 Word 文档内容。
        # - 提取文档中的段落和表格数据。
        #
        # ## 副作用/注意
        # - 只读操作，不会修改源文件。
        # - 复杂格式可能只保留纯文本。
        "description": """Read a .docx Word document.

## Prerequisites
python-docx must be installed. The path must use a namespace prefix (e.g. ws:, fork:).

## Effect
Returns a list of paragraphs (with text and style name), a list of tables (with headers and data rows), and a concatenated plain-text summary. Empty paragraphs are skipped.

## Returns
```json
{"path": "ws:docs/report.docx", "paragraphs": [{"style": "Normal", "text": "..."}], "tables": [{"headers": ["..."], "rows": [{"...": "..."}], "row_count": 1}], "plain_text": "...", "paragraph_count": 10, "table_count": 2}
```

## When to Use
- Read Word document content.
- Extract paragraphs and table data from a document.

## Side Effects / Notes
- Read-only operation; does not modify the source file.
- Complex formatting may be reduced to plain text.""",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    # Word 文档逻辑路径，必须使用命名空间前缀（ws:、fork:）。例如 'ws:docs/report.docx'。
                    "description": """Word document logical path, must use a namespace prefix (ws:, fork:). E.g. 'ws:docs/report.docx'.""",
                },
            },
            "required": ["path"],
        },
    },
    handler=_handle_read_docx,
    emoji="📄",
)